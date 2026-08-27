import os
import re
import json
import traceback
from datetime import datetime
from pathlib import Path

import joblib
import numpy as np

from flask import Flask, jsonify, render_template, request, session
from werkzeug.utils import secure_filename

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


# ============================================================
# FLASK CONFIGURATION
# ============================================================

BASE_DIR = Path(__file__).resolve().parent
MODEL_DIR = BASE_DIR / "models"
UPLOAD_DIR = BASE_DIR / "uploads"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get(
    "FLASK_SECRET_KEY",
    "career-intelligence-development-secret-key"
)

app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}


# ============================================================
# CAREER SKILL DATABASE
# ============================================================
# This is intentionally kept inside app.py so the requested
# project structure remains unchanged.
#
# The database is designed to be easy to extend.
# Skill names are normalized internally before comparison.
# ============================================================

CAREER_SKILLS = {
    "Data Scientist": [
        "python",
        "sql",
        "statistics",
        "machine learning",
        "deep learning",
        "pandas",
        "numpy",
        "scikit-learn",
        "tensorflow",
        "pytorch",
        "data visualization",
    ],

    "Machine Learning Engineer": [
        "python",
        "machine learning",
        "deep learning",
        "scikit-learn",
        "tensorflow",
        "pytorch",
        "numpy",
        "pandas",
        "sql",
        "statistics",
        "model deployment",
        "mlops",
    ],

    "Data Analyst": [
        "python",
        "sql",
        "excel",
        "statistics",
        "pandas",
        "numpy",
        "data visualization",
        "power bi",
        "tableau",
        "data analysis",
    ],

    "AI Engineer": [
        "python",
        "machine learning",
        "deep learning",
        "tensorflow",
        "pytorch",
        "scikit-learn",
        "numpy",
        "pandas",
        "natural language processing",
        "computer vision",
        "model deployment",
    ],

    "Software Engineer": [
        "python",
        "java",
        "javascript",
        "data structures",
        "algorithms",
        "object oriented programming",
        "git",
        "sql",
        "api",
        "testing",
    ],

    "Full Stack Developer": [
        "html",
        "css",
        "javascript",
        "python",
        "java",
        "react",
        "node.js",
        "sql",
        "git",
        "rest api",
        "web development",
    ],

    "Frontend Developer": [
        "html",
        "css",
        "javascript",
        "typescript",
        "react",
        "vue",
        "angular",
        "responsive design",
        "git",
        "web development",
    ],

    "Backend Developer": [
        "python",
        "java",
        "node.js",
        "sql",
        "rest api",
        "api",
        "databases",
        "git",
        "authentication",
        "server development",
    ],

    "Web Developer": [
        "html",
        "css",
        "javascript",
        "python",
        "php",
        "sql",
        "git",
        "responsive design",
        "web development",
    ],

    "DevOps Engineer": [
        "linux",
        "git",
        "docker",
        "kubernetes",
        "jenkins",
        "ci/cd",
        "aws",
        "azure",
        "cloud computing",
        "terraform",
        "python",
    ],

    "Cloud Engineer": [
        "aws",
        "azure",
        "google cloud",
        "cloud computing",
        "linux",
        "docker",
        "kubernetes",
        "terraform",
        "networking",
        "python",
    ],

    "Cybersecurity Analyst": [
        "cybersecurity",
        "network security",
        "linux",
        "penetration testing",
        "ethical hacking",
        "cryptography",
        "siem",
        "firewalls",
        "incident response",
        "risk assessment",
    ],

    "Database Administrator": [
        "sql",
        "mysql",
        "postgresql",
        "oracle",
        "database administration",
        "database design",
        "backup",
        "recovery",
        "performance tuning",
        "linux",
    ],

    "Business Analyst": [
        "business analysis",
        "sql",
        "excel",
        "power bi",
        "tableau",
        "requirements analysis",
        "data analysis",
        "communication",
        "problem solving",
    ],

    "Product Manager": [
        "product management",
        "product strategy",
        "market research",
        "agile",
        "scrum",
        "roadmap",
        "user research",
        "analytics",
        "communication",
    ],

    "UI/UX Designer": [
        "ui design",
        "ux design",
        "figma",
        "wireframing",
        "prototyping",
        "user research",
        "usability",
        "interaction design",
        "visual design",
    ],

    "Mobile App Developer": [
        "android",
        "ios",
        "java",
        "kotlin",
        "swift",
        "flutter",
        "react native",
        "mobile development",
        "git",
        "api",
    ],

    "Computer Vision Engineer": [
        "python",
        "computer vision",
        "opencv",
        "deep learning",
        "tensorflow",
        "pytorch",
        "numpy",
        "image processing",
        "cnn",
    ],

    "NLP Engineer": [
        "python",
        "natural language processing",
        "nlp",
        "machine learning",
        "deep learning",
        "transformers",
        "tensorflow",
        "pytorch",
        "text processing",
        "hugging face",
    ],

    "Data Engineer": [
        "python",
        "sql",
        "spark",
        "hadoop",
        "etl",
        "data pipelines",
        "airflow",
        "aws",
        "databases",
        "data warehousing",
    ],

    "Software Developer": [
        "python",
        "java",
        "javascript",
        "data structures",
        "algorithms",
        "git",
        "sql",
        "object oriented programming",
        "testing",
    ],

    "QA Engineer": [
        "software testing",
        "manual testing",
        "automation testing",
        "selenium",
        "python",
        "java",
        "api testing",
        "sql",
        "git",
        "test cases",
    ],
}


# ============================================================
# SKILL ALIASES
# ============================================================

SKILL_ALIASES = {
    "js": "javascript",
    "javascript": "javascript",
    "py": "python",
    "python3": "python",
    "python 3": "python",

    "ml": "machine learning",
    "machinelearning": "machine learning",

    "dl": "deep learning",
    "deeplearning": "deep learning",

    "sklearn": "scikit-learn",
    "scikit learn": "scikit-learn",

    "postgres": "postgresql",
    "postgres db": "postgresql",

    "node": "node.js",
    "nodejs": "node.js",

    "reactjs": "react",
    "react.js": "react",

    "vuejs": "vue",
    "vue.js": "vue",

    "angularjs": "angular",

    "powerbi": "power bi",

    "tableau software": "tableau",

    "nlp": "natural language processing",

    "cv": "computer vision",

    "rest": "rest api",
    "restful api": "rest api",

    "ci cd": "ci/cd",
    "cicd": "ci/cd",

    "aws cloud": "aws",
    "amazon web services": "aws",

    "gcp": "google cloud",
    "google cloud platform": "google cloud",

    "k8s": "kubernetes",

    "postgresql database": "postgresql",
}


# ============================================================
# LEARNING RECOMMENDATIONS
# ============================================================

SKILL_RECOMMENDATIONS = {
    "python": (
        "Strengthen Python fundamentals, object-oriented programming, "
        "data structures, libraries and practical project development."
    ),

    "sql": (
        "Practice joins, subqueries, CTEs, window functions, "
        "aggregation and database optimization."
    ),

    "machine learning": (
        "Study supervised and unsupervised learning, feature engineering, "
        "model evaluation, cross-validation and hyperparameter tuning."
    ),

    "deep learning": (
        "Learn neural networks, CNNs, RNNs, transformers and "
        "deep-learning model training using TensorFlow or PyTorch."
    ),

    "pandas": (
        "Practice DataFrame operations, filtering, grouping, merging, "
        "data cleaning and feature preparation."
    ),

    "numpy": (
        "Practice arrays, vectorized operations, matrix operations, "
        "broadcasting and numerical computation."
    ),

    "scikit-learn": (
        "Build classification, regression and clustering projects "
        "using preprocessing, pipelines and model evaluation."
    ),

    "tensorflow": (
        "Build neural-network projects and practice model training, "
        "validation, optimization and deployment."
    ),

    "pytorch": (
        "Practice tensors, neural networks, training loops, "
        "optimization and deep-learning projects."
    ),

    "statistics": (
        "Study probability, distributions, hypothesis testing, "
        "correlation, regression and statistical inference."
    ),

    "data visualization": (
        "Practice dashboards and visual storytelling using charts, "
        "matplotlib, Power BI or Tableau."
    ),

    "power bi": (
        "Build dashboards using Power Query, data modeling, DAX "
        "and interactive business reports."
    ),

    "tableau": (
        "Create interactive dashboards, calculated fields, "
        "filters and business-oriented visualizations."
    ),

    "excel": (
        "Practice formulas, pivot tables, lookup functions, "
        "charts and data analysis."
    ),

    "docker": (
        "Learn container images, Dockerfiles, networking, volumes "
        "and application containerization."
    ),

    "kubernetes": (
        "Learn pods, deployments, services, configuration, scaling "
        "and Kubernetes application management."
    ),

    "aws": (
        "Learn core AWS services such as EC2, S3, IAM, RDS, "
        "Lambda and cloud deployment."
    ),

    "azure": (
        "Learn Azure compute, storage, networking, identity "
        "and application deployment services."
    ),

    "linux": (
        "Practice Linux commands, permissions, processes, networking "
        "and shell-based system administration."
    ),

    "git": (
        "Practice branching, merging, pull requests, conflict resolution "
        "and professional Git workflows."
    ),

    "javascript": (
        "Practice modern JavaScript, DOM manipulation, asynchronous "
        "programming, APIs and modular application development."
    ),

    "html": (
        "Practice semantic HTML5, forms, accessibility and "
        "well-structured web page development."
    ),

    "css": (
        "Practice responsive layouts, Flexbox, Grid, animations "
        "and modern responsive design."
    ),

    "react": (
        "Build component-based applications using React, state management, "
        "routing and API integration."
    ),

    "node.js": (
        "Practice server-side JavaScript, Express-style APIs, "
        "authentication and database integration."
    ),

    "java": (
        "Strengthen Java fundamentals, OOP, collections, exception "
        "handling and backend development."
    ),

    "data structures": (
        "Practice arrays, linked lists, stacks, queues, trees, graphs "
        "and hash-based structures."
    ),

    "algorithms": (
        "Study sorting, searching, recursion, greedy methods, "
        "dynamic programming and graph algorithms."
    ),

    "natural language processing": (
        "Study tokenization, embeddings, transformers, text classification "
        "and language-model applications."
    ),

    "computer vision": (
        "Practice image preprocessing, feature extraction, CNNs, "
        "object detection and image classification."
    ),

    "spark": (
        "Learn distributed data processing, Spark DataFrames, "
        "Spark SQL and scalable ETL pipelines."
    ),

    "etl": (
        "Practice extracting, transforming and loading data using "
        "repeatable and reliable data pipelines."
    ),

    "ci/cd": (
        "Learn automated testing, build pipelines, deployment automation "
        "and continuous integration workflows."
    ),

    "api": (
        "Practice REST API design, HTTP methods, JSON, validation, "
        "authentication and API testing."
    ),
}


# ============================================================
# GLOBAL MODEL CONTAINER
# ============================================================

MODELS = {
    "logistic_regression": None,
    "random_forest": None,
    "xgboost": None,
}

VECTORIZERS = {
    "logistic_regression": None,
    "random_forest": None,
    "xgboost": None,
}

ENCODERS = {
    "logistic_regression": None,
    "random_forest": None,
    "xgboost": None,
}

MODEL_STATUS = {}


# ============================================================
# GENERAL HELPERS
# ============================================================

def model_path(filename):
    return MODEL_DIR / filename


def allowed_file(filename):
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
    )


def safe_float(value, default=0.0):
    try:
        value = float(value)

        if not np.isfinite(value):
            return default

        return value
    except Exception:
        return default


def normalize_text(text):
    if not text:
        return ""

    text = str(text).lower()

    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def normalize_skill(skill):
    skill = normalize_text(skill)

    if skill in SKILL_ALIASES:
        return SKILL_ALIASES[skill]

    return skill


def canonical_display_skill(skill):
    mapping = {
        "python": "Python",
        "sql": "SQL",
        "machine learning": "Machine Learning",
        "deep learning": "Deep Learning",
        "pandas": "Pandas",
        "numpy": "NumPy",
        "scikit-learn": "Scikit-learn",
        "tensorflow": "TensorFlow",
        "pytorch": "PyTorch",
        "statistics": "Statistics",
        "data visualization": "Data Visualization",
        "javascript": "JavaScript",
        "html": "HTML",
        "css": "CSS",
        "react": "React",
        "node.js": "Node.js",
        "java": "Java",
        "docker": "Docker",
        "kubernetes": "Kubernetes",
        "aws": "AWS",
        "azure": "Azure",
        "linux": "Linux",
        "git": "Git",
        "postgresql": "PostgreSQL",
        "mysql": "MySQL",
        "power bi": "Power BI",
        "tableau": "Tableau",
        "excel": "Excel",
        "natural language processing": "Natural Language Processing",
        "computer vision": "Computer Vision",
        "rest api": "REST API",
        "ci/cd": "CI/CD",
        "spark": "Apache Spark",
        "etl": "ETL",
        "data structures": "Data Structures",
        "algorithms": "Algorithms",
    }

    return mapping.get(skill, skill.title())


# ============================================================
# MODEL LOADING
# ============================================================

def load_pickle(filename):
    path = model_path(filename)

    if not path.exists():
        return None

    try:
        return joblib.load(path)
    except Exception as exc:
        print(f"[MODEL ERROR] Could not load {filename}: {exc}")
        return None


def load_models():
    """
    Load all available trained artifacts once when Flask starts.

    Model/vectorizer combinations:
        Logistic Regression -> tfidf.pkl + label_encoder.pkl
        Random Forest       -> rf_tfidf.pkl + rf_label_encoder.pkl
        XGBoost              -> tfidf.pkl + label_encoder.pkl

    If a model contains its own vectorizer/encoder, the application
    can also discover those attributes dynamically.
    """

    MODELS["logistic_regression"] = load_pickle(
        "logistic_regression.pkl"
    )

    MODELS["random_forest"] = load_pickle(
        "random_forest.pkl"
    )

    MODELS["xgboost"] = load_pickle(
        "xgboost.pkl"
    )

    common_tfidf = load_pickle("tfidf.pkl")
    common_encoder = load_pickle("label_encoder.pkl")

    rf_tfidf = load_pickle("rf_tfidf.pkl")
    rf_encoder = load_pickle("rf_label_encoder.pkl")

    VECTORIZERS["logistic_regression"] = common_tfidf
    VECTORIZERS["xgboost"] = common_tfidf
    VECTORIZERS["random_forest"] = rf_tfidf or common_tfidf

    ENCODERS["logistic_regression"] = common_encoder
    ENCODERS["xgboost"] = common_encoder
    ENCODERS["random_forest"] = rf_encoder or common_encoder

    for model_name, model in MODELS.items():
        if model is None:
            MODEL_STATUS[model_name] = {
                "available": False,
                "message": "Model file unavailable or failed to load."
            }
            continue

        # Try to discover embedded vectorizer.
        embedded_vectorizer = None

        for attr in [
            "vectorizer",
            "tfidf",
            "tfidf_vectorizer",
            "transformer",
        ]:
            if hasattr(model, attr):
                candidate = getattr(model, attr)

                if hasattr(candidate, "transform"):
                    embedded_vectorizer = candidate
                    break

        if embedded_vectorizer is not None:
            VECTORIZERS[model_name] = embedded_vectorizer

        # Try to discover embedded encoder.
        embedded_encoder = None

        for attr in [
            "label_encoder",
            "encoder",
        ]:
            if hasattr(model, attr):
                candidate = getattr(model, attr)

                if hasattr(candidate, "inverse_transform"):
                    embedded_encoder = candidate
                    break

        if embedded_encoder is not None:
            ENCODERS[model_name] = embedded_encoder

        MODEL_STATUS[model_name] = {
            "available": True,
            "message": "Model loaded successfully."
        }

    print("\n========== MODEL STATUS ==========")

    for name, status in MODEL_STATUS.items():
        print(
            f"{name}: "
            f"{'AVAILABLE' if status['available'] else 'UNAVAILABLE'}"
        )

    print("==================================\n")


# Load once at startup.
load_models()


# ============================================================
# CAREER LABEL HANDLING
# ============================================================

def get_encoder_classes(encoder):
    if encoder is None:
        return []

    try:
        classes = encoder.classes_

        return [
            str(item)
            for item in classes
        ]

    except Exception:
        return []


def get_model_classes(model_name):
    model = MODELS.get(model_name)
    encoder = ENCODERS.get(model_name)

    encoder_classes = get_encoder_classes(encoder)

    if encoder_classes:
        return encoder_classes

    if model is not None:
        try:
            classes = model.classes_

            return [
                str(item)
                for item in classes
            ]
        except Exception:
            pass

    return []


def get_all_careers():
    """
    Build a unified career list from the actual trained model encoders.
    This avoids hard-coding prediction labels.
    """

    careers = []

    for model_name in MODELS:
        for career in get_model_classes(model_name):
            if career not in careers:
                careers.append(career)

    return careers


# ============================================================
# RESUME EXTRACTION
# ============================================================

def extract_pdf_text(file_path):
    if PdfReader is None:
        raise RuntimeError(
            "PDF support is unavailable. Install the pypdf package."
        )

    reader = PdfReader(str(file_path))

    pages = []

    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")

    text = "\n".join(pages)

    if not text.strip():
        raise ValueError(
            "The PDF was uploaded successfully, but no readable text "
            "could be extracted."
        )

    return text


def extract_docx_text(file_path):
    if Document is None:
        raise RuntimeError(
            "DOCX support is unavailable. Install python-docx."
        )

    document = Document(str(file_path))

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                row_text.append(cell.text)

            parts.append(" ".join(row_text))

    text = "\n".join(parts)

    if not text.strip():
        raise ValueError(
            "The DOCX file was uploaded successfully, but no readable "
            "text was found."
        )

    return text


def extract_txt_text(file_path):
    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:
        text = file.read()

    if not text.strip():
        raise ValueError("The TXT file is empty.")

    return text


def extract_resume_text(file_path):
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    if extension == ".docx":
        return extract_docx_text(file_path)

    if extension == ".txt":
        return extract_txt_text(file_path)

    raise ValueError(
        "Unsupported resume file type."
    )


# ============================================================
# SKILL EXTRACTION
# ============================================================

def extract_skills(text):
    """
    Detect skills from the submitted resume/profile.

    Detection is intentionally based on the supported skill vocabulary,
    aliases and phrase matching rather than inventing skills.
    """

    normalized = normalize_text(text)

    if not normalized:
        return []

    detected = set()

    # Create a combined vocabulary from:
    # 1. career skill database
    # 2. skill aliases
    # 3. recommendation database

    vocabulary = set(SKILL_RECOMMENDATIONS.keys())

    for skills in CAREER_SKILLS.values():
        vocabulary.update(skills)

    vocabulary.update(SKILL_ALIASES.keys())

    # Longest phrases first prevents short aliases from interfering.
    vocabulary = sorted(
        vocabulary,
        key=len,
        reverse=True
    )

    for skill in vocabulary:
        normalized_skill = normalize_skill(skill)

        if not normalized_skill:
            continue

        escaped = re.escape(skill.lower())

        pattern = rf"(?<![a-z0-9+#]){escaped}(?![a-z0-9+#])"

        if re.search(pattern, normalized):
            detected.add(normalized_skill)

    # Handle a few symbols/names that benefit from special matching.
    special_patterns = {
        "c++": r"\bc\+\+\b",
        "c#": r"\bc#\b",
        ".net": r"\.net\b",
        "node.js": r"\bnode\.js\b",
        "ci/cd": r"\bci\s*/\s*cd\b",
    }

    for skill, pattern in special_patterns.items():
        if re.search(pattern, normalized):
            detected.add(normalize_skill(skill))

    return sorted(
        detected,
        key=lambda value: value.lower()
    )


# ============================================================
# SKILL GAP ENGINE
# ============================================================

def get_career_skill_list(career):
    """
    Return the configured skill requirements for a career.

    Exact matching is preferred.
    A case-insensitive lookup is also supported.
    """

    if not career:
        return []

    if career in CAREER_SKILLS:
        return [
            normalize_skill(skill)
            for skill in CAREER_SKILLS[career]
        ]

    career_lower = normalize_text(career)

    for known_career, skills in CAREER_SKILLS.items():
        if normalize_text(known_career) == career_lower:
            return [
                normalize_skill(skill)
                for skill in skills
            ]

    return []


def calculate_skill_gap(candidate_skills, career):
    candidate = {
        normalize_skill(skill)
        for skill in candidate_skills
        if skill
    }

    required = get_career_skill_list(career)

    if not required:
        return {
            "available": False,
            "career": career,
            "required_skills": [],
            "candidate_skills": [
                canonical_display_skill(skill)
                for skill in sorted(candidate)
            ],
            "matched_skills": [],
            "missing_skills": [],
            "matched_count": 0,
            "missing_count": 0,
            "skill_match": 0.0,
            "skill_gap": 0.0,
            "readiness_score": None,
            "priority": {
                "high": [],
                "medium": [],
                "low": []
            }
        }

    required_set = set(required)

    matched = sorted(
        candidate.intersection(required_set)
    )

    missing = sorted(
        required_set.difference(candidate)
    )

    match_percentage = (
        len(matched) / len(required_set)
    ) * 100

    gap_percentage = 100 - match_percentage

    # Priority is derived from the order of the career's skill
    # requirements. The first third are treated as high priority,
    # the middle third as medium and the remaining as low.
    #
    # This is a maintainable deterministic rule rather than random
    # classification.
    high_count = max(
        1,
        int(np.ceil(len(required) * 0.35))
    )

    medium_count = max(
        high_count + 1,
        int(np.ceil(len(required) * 0.70))
    )

    priority = {
        "high": [],
        "medium": [],
        "low": []
    }

    for index, skill in enumerate(required):
        if skill not in missing:
            continue

        if index < high_count:
            priority["high"].append(skill)
        elif index < medium_count:
            priority["medium"].append(skill)
        else:
            priority["low"].append(skill)

    readiness_score = round(match_percentage)

    return {
        "available": True,
        "career": career,
        "required_skills": [
            canonical_display_skill(skill)
            for skill in required
        ],
        "candidate_skills": [
            canonical_display_skill(skill)
            for skill in sorted(candidate)
        ],
        "matched_skills": [
            canonical_display_skill(skill)
            for skill in matched
        ],
        "missing_skills": [
            canonical_display_skill(skill)
            for skill in missing
        ],
        "matched_count": len(matched),
        "missing_count": len(missing),
        "skill_match": round(match_percentage, 2),
        "skill_gap": round(gap_percentage, 2),
        "readiness_score": readiness_score,
        "priority": {
            "high": [
                canonical_display_skill(skill)
                for skill in priority["high"]
            ],
            "medium": [
                canonical_display_skill(skill)
                for skill in priority["medium"]
            ],
            "low": [
                canonical_display_skill(skill)
                for skill in priority["low"]
            ]
        }
    }


def build_recommendations(skill_gap):
    recommendations = []

    for priority_name in ["high", "medium", "low"]:
        for display_skill in skill_gap.get(
            "priority",
            {}
        ).get(priority_name, []):

            normalized = normalize_skill(display_skill)

            recommendation = SKILL_RECOMMENDATIONS.get(
                normalized,
                (
                    f"Develop practical {display_skill} knowledge "
                    f"through structured learning and a hands-on project."
                )
            )

            recommendations.append({
                "skill": display_skill,
                "priority": priority_name,
                "recommendation": recommendation
            })

    return recommendations


# ============================================================
# MODEL PREDICTION HELPERS
# ============================================================

def decode_prediction(model_name, prediction):
    encoder = ENCODERS.get(model_name)

    # Prediction may already be a career string.
    if isinstance(prediction, str):
        return prediction

    if encoder is not None:
        try:
            decoded = encoder.inverse_transform(
                np.asarray([prediction])
            )

            return str(decoded[0])
        except Exception:
            pass

    model_classes = get_model_classes(model_name)

    try:
        numeric_prediction = int(prediction)

        if (
            0 <= numeric_prediction
            < len(model_classes)
        ):
            return model_classes[numeric_prediction]
    except Exception:
        pass

    return str(prediction)


def get_probability_vector(model_name, model, features):
    """
    Return:
        classes, probabilities

    If predict_proba is unavailable, return a deterministic one-hot
    prediction. This is not a fabricated probability; it represents
    the model's hard prediction/vote.
    """

    if model is None:
        return [], np.array([])

    classes = get_model_classes(model_name)

    # Preferred path.
    if hasattr(model, "predict_proba"):
        try:
            probabilities = model.predict_proba(features)

            probabilities = np.asarray(
                probabilities,
                dtype=float
            )

            if probabilities.ndim == 2:
                probabilities = probabilities[0]

            if not classes:
                try:
                    raw_classes = model.classes_

                    classes = [
                        decode_prediction(
                            model_name,
                            item
                        )
                        for item in raw_classes
                    ]
                except Exception:
                    pass

            if len(classes) == len(probabilities):
                return classes, probabilities

        except Exception as exc:
            print(
                f"[PREDICT_PROBA WARNING] "
                f"{model_name}: {exc}"
            )

    # Safe hard-prediction fallback.
    try:
        prediction = model.predict(features)[0]

        career = decode_prediction(
            model_name,
            prediction
        )

        if not classes:
            classes = [career]

        probabilities = np.zeros(
            len(classes),
            dtype=float
        )

        if career in classes:
            probabilities[
                classes.index(career)
            ] = 1.0
        else:
            classes.append(career)

            probabilities = np.append(
                probabilities,
                1.0
            )

        return classes, probabilities

    except Exception as exc:
        print(
            f"[PREDICTION ERROR] "
            f"{model_name}: {exc}"
        )

        return [], np.array([])


def transform_for_model(model_name, resume_text):
    vectorizer = VECTORIZERS.get(model_name)

    if vectorizer is None:
        raise RuntimeError(
            f"No TF-IDF/vectorizer is available for "
            f"{model_name}."
        )

    return vectorizer.transform([resume_text])


def predict_single_model(model_name, resume_text):
    model = MODELS.get(model_name)

    if model is None:
        return {
            "available": False,
            "career": None,
            "confidence": None,
            "predictions": [],
            "error": "Model is unavailable."
        }

    try:
        features = transform_for_model(
            model_name,
            resume_text
        )

        classes, probabilities = get_probability_vector(
            model_name,
            model,
            features
        )

        if len(classes) == 0 or len(probabilities) == 0:
            raise RuntimeError(
                "Model returned no usable predictions."
            )

        # Sanitize numerical values.
        probabilities = np.asarray(
            probabilities,
            dtype=float
        )

        probabilities[
            ~np.isfinite(probabilities)
        ] = 0.0

        probabilities = np.clip(
            probabilities,
            0.0,
            None
        )

        total = probabilities.sum()

        if total > 0:
            probabilities = probabilities / total

        ranking = sorted(
            zip(classes, probabilities),
            key=lambda item: item[1],
            reverse=True
        )

        predictions = []

        for career, probability in ranking[:5]:
            predictions.append({
                "career": str(career),
                "confidence": round(
                    float(probability * 100),
                    2
                )
            })

        top_career = predictions[0]

        return {
            "available": True,
            "career": top_career["career"],
            "confidence": top_career["confidence"],
            "predictions": predictions,
            "error": None
        }

    except Exception as exc:
        print(
            f"[MODEL PREDICTION ERROR] "
            f"{model_name}: {exc}"
        )

        return {
            "available": False,
            "career": None,
            "confidence": None,
            "predictions": [],
            "error": str(exc)
        }


# ============================================================
# ENSEMBLE PREDICTION
# ============================================================

def ensemble_predictions(
    model_results,
    candidate_skills
):
    """
    Ensemble strategy:

        Logistic Regression = 1/3
        Random Forest       = 1/3
        XGBoost              = 1/3

    Only available models contribute.

    Each model's probability distribution is normalized and aligned
    by career name.

    When a model does not expose predict_proba(), its hard prediction
    becomes a deterministic one-vote distribution. No random values
    are introduced.
    """

    usable = []

    for model_name, result in model_results.items():
        if result.get("available"):
            usable.append(
                (model_name, result)
            )

    if not usable:
        return {
            "top_career": None,
            "predictions": [],
            "ensemble_method": "No usable models",
        }

    career_scores = {}

    for model_name, result in usable:
        predictions = result.get(
            "predictions",
            []
        )

        for item in predictions:
            career = item["career"]

            confidence = safe_float(
                item.get("confidence"),
                0
            )

            career_scores.setdefault(
                career,
                []
            ).append(
                confidence
            )

    # Average available model confidence for each career.
    ensemble_rows = []

    for career, values in career_scores.items():
        if not values:
            continue

        ensemble_confidence = sum(values) / len(values)

        skill_gap = calculate_skill_gap(
            candidate_skills,
            career
        )

        skill_match = skill_gap.get(
            "skill_match",
            0.0
        )

        # Ranking combines:
        # 70% ML ensemble confidence
        # 30% skill alignment
        #
        # This prevents a high-confidence model prediction from
        # completely ignoring the candidate's detected skills.
        ranking_score = (
            ensemble_confidence * 0.70
            + skill_match * 0.30
        )

        ensemble_rows.append({
            "career": career,
            "confidence": round(
                ensemble_confidence,
                2
            ),
            "skill_match": round(
                skill_match,
                2
            ),
            "skill_gap": round(
                skill_gap.get(
                    "skill_gap",
                    100.0
                ),
                2
            ),
            "ranking_score": round(
                ranking_score,
                2
            ),
        })

    ensemble_rows.sort(
        key=lambda item: item["ranking_score"],
        reverse=True
    )

    top_five = ensemble_rows[:5]

    for index, row in enumerate(top_five, start=1):
        row["rank"] = index

    return {
        "top_career": (
            top_five[0]
            if top_five
            else None
        ),
        "predictions": top_five,
        "ensemble_method": (
            "Average available model confidence "
            "with 70% ML confidence + 30% skill alignment ranking"
        )
    }


# ============================================================
# BUILD MODEL RESPONSE
# ============================================================

def build_model_response(model_results):
    output = {}

    for model_name, result in model_results.items():
        output[model_name] = {
            "available": result.get(
                "available",
                False
            ),
            "career": result.get(
                "career"
            ),
            "confidence": result.get(
                "confidence"
            ),
            "predictions": result.get(
                "predictions",
                []
            ),
            "error": result.get(
                "error"
            )
        }

    return output


# ============================================================
# COMPLETE ANALYSIS
# ============================================================

def analyze_profile(resume_text):
    normalized_resume = normalize_text(
        resume_text
    )

    if not normalized_resume:
        raise ValueError(
            "Resume or skills text cannot be empty."
        )

    if len(normalized_resume) < 3:
        raise ValueError(
            "Please provide more resume or skill information."
        )

    candidate_skills = extract_skills(
        normalized_resume
    )

    model_results = {}

    for model_name in [
        "logistic_regression",
        "random_forest",
        "xgboost"
    ]:
        model_results[model_name] = predict_single_model(
            model_name,
            normalized_resume
        )

    ensemble = ensemble_predictions(
        model_results,
        candidate_skills
    )

    top_career = ensemble.get(
        "top_career"
    )

    if top_career is None:
        errors = []

        for name, result in model_results.items():
            if result.get("error"):
                errors.append(
                    f"{name}: {result['error']}"
                )

        raise RuntimeError(
            "No trained model could produce a prediction. "
            + (
                " | ".join(errors)
                if errors
                else "Check your model files and vectorizers."
            )
        )

    top_career_name = top_career["career"]

    skill_gap = calculate_skill_gap(
        candidate_skills,
        top_career_name
    )

    recommendations = build_recommendations(
        skill_gap
    )

    predictions = []

    for prediction in ensemble["predictions"]:
        career_name = prediction["career"]

        career_gap = calculate_skill_gap(
            candidate_skills,
            career_name
        )

        predictions.append({
            "rank": prediction["rank"],
            "career": career_name,
            "confidence": prediction["confidence"],
            "skill_match": career_gap.get(
                "skill_match",
                0.0
            ),
            "skill_gap": career_gap.get(
                "skill_gap",
                100.0
            ),
            "matched_count": career_gap.get(
                "matched_count",
                0
            ),
            "missing_count": career_gap.get(
                "missing_count",
                0
            )
        })

    top_prediction = predictions[0]

    skill_gap["recommendations"] = recommendations

    response = {
        "success": True,

        "timestamp": datetime.now().isoformat(),

        "top_career": {
            "name": top_prediction["career"],
            "confidence": top_prediction["confidence"],
            "skill_match": skill_gap.get(
                "skill_match",
                0.0
            ),
            "skill_gap": skill_gap.get(
                "skill_gap",
                100.0
            ),
            "readiness_score": skill_gap.get(
                "readiness_score"
            ),
            "description": (
                f"{top_prediction['career']} is the highest-ranked "
                "career based on the available machine-learning "
                "predictions and your detected skill alignment."
            ),
            "matched_count": skill_gap.get(
                "matched_count",
                0
            ),
            "missing_count": skill_gap.get(
                "missing_count",
                0
            )
        },

        "predictions": predictions,

        "models": build_model_response(
            model_results
        ),

        "skills": {
            "identified": [
                canonical_display_skill(skill)
                for skill in candidate_skills
            ],
            "matched": skill_gap.get(
                "matched_skills",
                []
            ),
            "missing": skill_gap.get(
                "missing_skills",
                []
            ),
            "required": skill_gap.get(
                "required_skills",
                []
            )
        },

        "skill_gap": skill_gap,

        "recommendations": recommendations,

        "ensemble": {
            "method": ensemble.get(
                "ensemble_method"
            )
        },

        "model_status": MODEL_STATUS
    }

    return response


# ============================================================
# LOGIN API
# ============================================================

@app.post("/login")
def login():
    data = request.get_json(
        silent=True
    ) or {}

    email = str(
        data.get("email", "")
    ).strip()

    password = str(
        data.get("password", "")
    )

    if not email or not password:
        return jsonify({
            "success": False,
            "message": "Email and password are required."
        }), 400

    # Demo authentication requested by the specification.
    if (
        email.lower() == "demo@example.com"
        and password == "demo123"
    ):
        session["logged_in"] = True
        session["user_email"] = email

        return jsonify({
            "success": True,
            "message": "Login successful.",
            "user": {
                "email": email,
                "name": "Demo User"
            }
        })

    return jsonify({
        "success": False,
        "message": (
            "Invalid credentials. Use "
            "demo@example.com / demo123."
        )
    }), 401


@app.post("/logout")
def logout():
    session.clear()

    return jsonify({
        "success": True,
        "message": "Logged out successfully."
    })


# ============================================================
# AUTH HELPER
# ============================================================

def is_authenticated():
    return bool(
        session.get("logged_in")
    )


# ============================================================
# ANALYZE ENDPOINT
# ============================================================

@app.post("/analyze")
def analyze():
    if not is_authenticated():
        return jsonify({
            "success": False,
            "message": "Please log in before analyzing a resume."
        }), 401

    resume_text = ""
    saved_file = None

    try:
        # ----------------------------------------------------
        # JSON request
        # ----------------------------------------------------
        if request.is_json:
            data = request.get_json(
                silent=True
            ) or {}

            resume_text = str(
                data.get(
                    "resume_text",
                    ""
                )
            )

        # ----------------------------------------------------
        # Multipart request
        # ----------------------------------------------------
        else:
            resume_text = str(
                request.form.get(
                    "resume_text",
                    ""
                )
            )

            uploaded_file = request.files.get(
                "resume"
            )

            if uploaded_file and uploaded_file.filename:
                original_filename = (
                    uploaded_file.filename
                )

                if not allowed_file(
                    original_filename
                ):
                    return jsonify({
                        "success": False,
                        "message": (
                            "Unsupported file type. "
                            "Please upload PDF, DOCX or TXT."
                        )
                    }), 400

                filename = secure_filename(
                    original_filename
                )

                if not filename:
                    return jsonify({
                        "success": False,
                        "message": "Invalid filename."
                    }), 400

                timestamp = datetime.now().strftime(
                    "%Y%m%d_%H%M%S_%f"
                )

                stored_filename = (
                    f"{timestamp}_{filename}"
                )

                saved_file = (
                    UPLOAD_DIR
                    / stored_filename
                )

                uploaded_file.save(
                    str(saved_file)
                )

                extracted_text = extract_resume_text(
                    saved_file
                )

                # Uploaded resume takes precedence.
                resume_text = extracted_text

        if not resume_text.strip():
            return jsonify({
                "success": False,
                "message": (
                    "Please upload a resume or enter your "
                    "skills/resume text."
                )
            }), 400

        if len(resume_text) > 500000:
            return jsonify({
                "success": False,
                "message": (
                    "The submitted text is too large. "
                    "Please provide a concise resume."
                )
            }), 400

        result = analyze_profile(
            resume_text
        )

        # Store only lightweight information in Flask session.
        session["last_career"] = result[
            "top_career"
        ]["name"]

        return jsonify(result)

    except ValueError as exc:
        return jsonify({
            "success": False,
            "message": str(exc)
        }), 400

    except RuntimeError as exc:
        return jsonify({
            "success": False,
            "message": str(exc)
        }), 500

    except Exception as exc:
        print(
            "[ANALYZE ERROR]"
        )

        traceback.print_exc()

        return jsonify({
            "success": False,
            "message": (
                "An unexpected error occurred while "
                "analyzing the profile."
            )
        }), 500

    finally:
        # The uploaded resume is temporary.
        # Do not expose it through Flask.
        if saved_file and saved_file.exists():
            try:
                saved_file.unlink()
            except Exception:
                pass


# ============================================================
# CAREER DETAILS API
# ============================================================

@app.post("/career-details")
def career_details():
    if not is_authenticated():
        return jsonify({
            "success": False,
            "message": "Please log in."
        }), 401

    data = request.get_json(
        silent=True
    ) or {}

    career = str(
        data.get("career", "")
    ).strip()

    if not career:
        return jsonify({
            "success": False,
            "message": "Career name is required."
        }), 400

    candidate_skills = data.get(
        "candidate_skills",
        []
    )

    if not isinstance(
        candidate_skills,
        list
    ):
        candidate_skills = []

    gap = calculate_skill_gap(
        candidate_skills,
        career
    )

    recommendations = build_recommendations(
        gap
    )

    return jsonify({
        "success": True,
        "career": career,
        "skill_gap": gap,
        "recommendations": recommendations
    })


# ============================================================
# SKILL GAP API
# ============================================================

@app.post("/skill-gap")
def skill_gap():
    if not is_authenticated():
        return jsonify({
            "success": False,
            "message": "Please log in."
        }), 401

    data = request.get_json(
        silent=True
    ) or {}

    career = str(
        data.get("career", "")
    ).strip()

    candidate_skills = data.get(
        "candidate_skills",
        []
    )

    if not career:
        return jsonify({
            "success": False,
            "message": "Career is required."
        }), 400

    if not isinstance(
        candidate_skills,
        list
    ):
        return jsonify({
            "success": False,
            "message": "candidate_skills must be a list."
        }), 400

    result = calculate_skill_gap(
        candidate_skills,
        career
    )

    result["recommendations"] = build_recommendations(
        result
    )

    return jsonify({
        "success": True,
        **result
    })


# ============================================================
# PREDICT API
# ============================================================

@app.post("/predict")
def predict():
    if not is_authenticated():
        return jsonify({
            "success": False,
            "message": "Please log in."
        }), 401

    data = request.get_json(
        silent=True
    ) or {}

    resume_text = str(
        data.get(
            "resume_text",
            ""
        )
    ).strip()

    if not resume_text:
        return jsonify({
            "success": False,
            "message": "resume_text is required."
        }), 400

    try:
        result = analyze_profile(
            resume_text
        )

        return jsonify({
            "success": True,
            "top_career": result["top_career"],
            "predictions": result["predictions"],
            "models": result["models"]
        })

    except Exception as exc:
        return jsonify({
            "success": False,
            "message": str(exc)
        }), 500


# ============================================================
# HEALTH API
# ============================================================

@app.get("/health")
def health():
    return jsonify({
        "status": "ok",
        "application": "AI Powered Career Intelligence Platform",
        "models": MODEL_STATUS,
        "career_count": len(
            get_all_careers()
        ),
        "timestamp": datetime.now().isoformat()
    })


# ============================================================
# CONFIGURATION / MODEL INFO
# ============================================================

@app.get("/api/model-status")
def model_status():
    return jsonify({
        "success": True,
        "models": MODEL_STATUS,
        "careers": get_all_careers()
    })


# ============================================================
# MAIN PAGE
# ============================================================

@app.get("/")
def index():
    return render_template(
        "index.html"
    )


# ============================================================
# ERROR HANDLERS
# ============================================================

@app.errorhandler(413)
def request_too_large(error):
    return jsonify({
        "success": False,
        "message": (
            "Uploaded file is too large. "
            "Maximum allowed size is 10 MB."
        )
    }), 413


@app.errorhandler(404)
def not_found(error):
    if request.path.startswith("/api") or request.path in {
        "/analyze",
        "/predict",
        "/skill-gap",
        "/career-details",
        "/login",
        "/logout"
    }:
        return jsonify({
            "success": False,
            "message": "Endpoint not found."
        }), 404

    return render_template(
        "index.html"
    )


# ============================================================
# DEVELOPMENT SERVER
# ============================================================

if __name__ == "__main__":
    print(
        "\nAI Powered Career Intelligence Platform"
    )
    print(
        "Open: http://127.0.0.1:5000"
    )

    app.run(
        host="127.0.0.1",
        port=5000,
        debug=True
    )